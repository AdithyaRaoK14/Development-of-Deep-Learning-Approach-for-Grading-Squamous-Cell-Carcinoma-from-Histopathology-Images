"""
=======================================================================
PROJECT: Lung Cancer Classification (ResNet152) - MULTI-EXPERIMENT
AUTHOR: Modified for multiple experiments with different configurations
CHANGES:
  - Added experiment configuration system
  - Multiple learning rate experiments (EXP_1 to EXP_6)
  - Dropout experiments (EXP_7 to EXP_10)
  - Regularization experiments (EXP_11 to EXP_14)
  - Fixed Grad-CAM OpenCV error
  - Experiment ID integrated into output paths
  - COMPLETE: All report generation (fold + comprehensive) with PDF
  - VERIFIED: All sections present and functional
=======================================================================
"""

# ================== IMPORTS ==================
import os
import random
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
import tensorflow as tf
from tensorflow.keras.preprocessing.image import ImageDataGenerator
from tensorflow.keras.applications import ResNet152
from tensorflow.keras.models import Model
from tensorflow.keras.layers import Dense, GlobalAveragePooling2D, Dropout
from tensorflow.keras.optimizers import AdamW
from tensorflow.keras.callbacks import ReduceLROnPlateau, EarlyStopping, CSVLogger
from tensorflow.keras import regularizers
from sklearn.model_selection import KFold
from sklearn.metrics import (confusion_matrix, ConfusionMatrixDisplay, classification_report, 
                             balanced_accuracy_score, roc_curve, auc, roc_auc_score)
from sklearn.preprocessing import label_binarize
import cv2

# Try to import docx for report generation
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    DOCX_AVAILABLE = True
    print(" python-docx available - Word reports will be generated")
except ImportError:
    DOCX_AVAILABLE = False
    print(" python-docx not installed. Install with: pip install python-docx")
    print("   Reports will be skipped.")

# Try to import docx2pdf for PDF conversion
try:
    from docx2pdf import convert
    PDF_AVAILABLE = True
    print(" docx2pdf available - PDF reports will be generated")
except ImportError:
    PDF_AVAILABLE = False
    print(" docx2pdf not installed. Install with: pip install docx2pdf")
    print("   PDF conversion will be skipped.")

sns.set(style="whitegrid")

# ============================================================
#  EXPERIMENT CONFIGURATIONS
# ============================================================
EXP_ID = "EXP_10"  # <--- CHANGE THIS TO RUN DIFFERENT EXPERIMENTS

EXPERIMENTS = {
    # Learning Rate Experiments
    "EXP_1": {
        "name": "Baseline_LR_1e-4",
        "learning_rate": 1e-4,
        "min_lr": 1e-7,
        "dropout": None,
        "regularization": None,
        "description": "Baseline with LR=1e-4"
    },
    "EXP_2": {
        "name": "High_LR_5e-4",
        "learning_rate": 5e-4,
        "min_lr": 1e-7,
        "dropout": None,
        "regularization": None,
        "description": "Higher learning rate (5x baseline)"
    },
    "EXP_3": {
        "name": "Low_LR_5e-5",
        "learning_rate": 5e-5,
        "min_lr": 1e-7,
        "dropout": None,
        "regularization": None,
        "description": "Lower learning rate (0.5x baseline)"
    },
    "EXP_4": {
        "name": "VeryLow_LR_1e-5",
        "learning_rate": 1e-5,
        "min_lr": 1e-7,
        "dropout": None,
        "regularization": None,
        "description": "Very low learning rate (0.1x baseline)"
    },
    "EXP_5": {
        "name": "VeryHigh_LR_1e-3",
        "learning_rate": 1e-3,
        "min_lr": 1e-7,
        "dropout": None,
        "regularization": None,
        "description": "Very high learning rate (10x baseline)"
    },
    "EXP_6": {
        "name": "UltraLow_LR_5e-6",
        "learning_rate": 5e-6,
        "min_lr": 1e-7,
        "dropout": None,
        "regularization": None,
        "description": "Ultra low learning rate for fine-grained training"
    },
    
    # Dropout Experiments
    "EXP_7": {
        "name": "Dropout_0.2_LR_1e-4",
        "learning_rate": 1e-4,
        "min_lr": 1e-7,
        "dropout": 0.2,
        "regularization": None,
        "description": "Light dropout (20%) with baseline LR"
    },
    "EXP_8": {
        "name": "Dropout_0.3_LR_1e-4",
        "learning_rate": 1e-4,
        "min_lr": 1e-7,
        "dropout": 0.3,
        "regularization": None,
        "description": "Moderate dropout (30%) with baseline LR"
    },
    "EXP_9": {
        "name": "Dropout_0.5_LR_1e-4",
        "learning_rate": 1e-4,
        "min_lr": 1e-7,
        "dropout": 0.5,
        "regularization": None,
        "description": "Heavy dropout (50%) with baseline LR"
    },
    "EXP_10": {
        "name": "Dropout_0.4_LR_5e-5",
        "learning_rate": 5e-5,
        "min_lr": 1e-7,
        "dropout": 0.4,
        "regularization": None,
        "description": "Dropout (40%) with lower LR"
    },
    
    # Regularization Experiments
    "EXP_11": {
        "name": "L2_1e-4_LR_1e-4",
        "learning_rate": 1e-4,
        "min_lr": 1e-7,
        "dropout": None,
        "regularization": {"type": "l2", "value": 1e-4},
        "description": "L2 regularization (1e-4) with baseline LR"
    },
    "EXP_12": {
        "name": "L2_1e-3_LR_1e-4",
        "learning_rate": 1e-4,
        "min_lr": 1e-7,
        "dropout": None,
        "regularization": {"type": "l2", "value": 1e-3},
        "description": "Strong L2 regularization (1e-3) with baseline LR"
    },
    "EXP_13": {
        "name": "L2_1e-4_Dropout_0.3_LR_1e-4",
        "learning_rate": 1e-4,
        "min_lr": 1e-7,
        "dropout": 0.3,
        "regularization": {"type": "l2", "value": 1e-4},
        "description": "Combined L2 regularization and dropout"
    },
    "EXP_14": {
        "name": "L1_1e-4_LR_1e-4",
        "learning_rate": 1e-4,
        "min_lr": 1e-7,
        "dropout": None,
        "regularization": {"type": "l1", "value": 1e-4},
        "description": "L1 regularization (1e-4) with baseline LR"
    },
}

# Validate experiment ID
if EXP_ID not in EXPERIMENTS:
    raise ValueError(f"Invalid EXP_ID: {EXP_ID}. Choose from {list(EXPERIMENTS.keys())}")

# Get current experiment config
exp_config = EXPERIMENTS[EXP_ID]
LEARNING_RATE = exp_config["learning_rate"]
MIN_LR = exp_config["min_lr"]
DROPOUT = exp_config["dropout"]
REGULARIZATION = exp_config["regularization"]

print(f"\n{'='*70}")
print(f"🔬 RUNNING EXPERIMENT: {EXP_ID}")
print(f"{'='*70}")
print(f" Name: {exp_config['name']}")
print(f" Description: {exp_config['description']}")
print(f" Learning Rate: {LEARNING_RATE}")
print(f" Min Learning Rate: {MIN_LR}")
print(f" Dropout: {DROPOUT if DROPOUT else 'None'}")
print(f" Regularization: {REGULARIZATION if REGULARIZATION else 'None'}")
print(f"{'='*70}\n")

# ============================================================
# 1️ PATH & SEED SETUP
# ============================================================
SEED = 42
np.random.seed(SEED)
random.seed(SEED)
tf.random.set_seed(SEED)
tf.keras.utils.set_random_seed(SEED)

base_path = "/data/vaishnav25/Data"
output_dir = f"/data/vaishnav25/ADITHYA/RESNET152_{EXP_ID}"
plot_dir = os.path.join(output_dir, "plots")
model_dir = os.path.join(output_dir, "models")
gradcam_dir = os.path.join(output_dir, "gradcams")
log_dir = os.path.join(output_dir, "logs")

for d in [plot_dir, model_dir, gradcam_dir, log_dir]:
    os.makedirs(d, exist_ok=True)

# Save experiment config
with open(os.path.join(output_dir, "experiment_config.txt"), "w") as f:
    f.write(f"Experiment ID: {EXP_ID}\n")
    f.write(f"Name: {exp_config['name']}\n")
    f.write(f"Description: {exp_config['description']}\n")
    f.write(f"Learning Rate: {LEARNING_RATE}\n")
    f.write(f"Min Learning Rate: {MIN_LR}\n")
    f.write(f"Dropout: {DROPOUT}\n")
    f.write(f"Regularization: {REGULARIZATION}\n")

class_names = ['Mod', 'Poor', 'Well']
class_paths = [os.path.join(base_path, c) for c in class_names]

# Hyperparameters
INPUT_SIZE = 224
BATCH_SIZE = 8
TARGET_IMAGES_PER_CLASS = 500

# ============================================================
# 2️ GPU CHECK & MEMORY CONTROL
# ============================================================
from tensorflow.keras import mixed_precision
policy = mixed_precision.Policy('mixed_float16')
mixed_precision.set_global_policy(policy)
print(f" Mixed precision enabled: {policy.name}")

gpus = tf.config.list_physical_devices('GPU')
if gpus:
    try:
        for gpu in gpus:
            tf.config.experimental.set_memory_growth(gpu, True)
        print(f" GPU available: {len(gpus)} device(s)")
        print(f"   Memory growth enabled")
        print(f"   Batch size: {BATCH_SIZE}")
    except RuntimeError as e:
        print(" GPU memory config error:", e)
else:
    print("⚙️ Using CPU (no GPU detected).")

# ============================================================
# 3️ LOAD AND BALANCE DATA
# ============================================================
print(f"\n{'='*60}")
print(f"📊 LOADING AND BALANCING DATASET")
print(f"{'='*60}")

all_image_paths = {cls: [] for cls in class_names}
for cls_name, class_dir in zip(class_names, class_paths):
    if not os.path.isdir(class_dir):
        raise FileNotFoundError(f"Missing directory: {class_dir}")
    for fname in os.listdir(class_dir):
        if fname.lower().endswith(('.png', '.jpg', '.jpeg', '.tif')):
            all_image_paths[cls_name].append(os.path.join(class_dir, fname))

print("\nOriginal dataset distribution:")
for cls in class_names:
    print(f"  {cls:6s}: {len(all_image_paths[cls])} images")

balanced_image_paths = []
balanced_labels = []

for cls_name, cls_idx in [('Mod', 0), ('Well', 2)]:
    images = all_image_paths[cls_name]
    selected = np.random.choice(images, TARGET_IMAGES_PER_CLASS, replace=False)
    balanced_image_paths.extend(selected)
    balanced_labels.extend([cls_idx] * TARGET_IMAGES_PER_CLASS)

poor_images = all_image_paths['Poor']
balanced_image_paths.extend(poor_images)
balanced_labels.extend([1] * len(poor_images))

print(f"\n Balanced dataset (before Poor augmentation):")
print(f"  Mod : {TARGET_IMAGES_PER_CLASS} images")
print(f"  Poor: {len(poor_images)} images (will augment to {TARGET_IMAGES_PER_CLASS})")
print(f"  Well: {TARGET_IMAGES_PER_CLASS} images")

# ============================================================
# 4️ AUGMENTATION FOR POOR CLASS
# ============================================================
color_augmentation = ImageDataGenerator(
    rescale=1./255,
    brightness_range=[0.8, 1.2],
    channel_shift_range=20.0,
    fill_mode='nearest'
)

print(f"\n Augmenting Poor class images...")
augmented_poor_dir = os.path.join(output_dir, "augmented_poor")
os.makedirs(augmented_poor_dir, exist_ok=True)

num_to_augment = TARGET_IMAGES_PER_CLASS - len(poor_images)
augmented_count = 0

while augmented_count < num_to_augment:
    source_img_path = np.random.choice(poor_images)
    img = tf.keras.preprocessing.image.load_img(source_img_path, target_size=(INPUT_SIZE, INPUT_SIZE))
    img_array = tf.keras.preprocessing.image.img_to_array(img)
    img_array = np.expand_dims(img_array, axis=0)
    
    aug_iter = color_augmentation.flow(img_array, batch_size=1)
    aug_img = next(aug_iter)[0]
    
    aug_img_path = os.path.join(augmented_poor_dir, f"aug_poor_{augmented_count}.png")
    tf.keras.preprocessing.image.save_img(aug_img_path, aug_img)
    
    balanced_image_paths.append(aug_img_path)
    balanced_labels.append(1)
    augmented_count += 1

print(f" Generated {augmented_count} augmented images for Poor class")

image_paths = np.array(balanced_image_paths)
labels = np.array(balanced_labels)

print(f"\n{'='*60}")
print(f" FINAL BALANCED DATASET STATISTICS")
print(f"{'='*60}")
print(f"Total images: {len(image_paths)}")
for i, c in enumerate(class_names):
    count = np.sum(labels==i)
    pct = 100 * count / len(image_paths)
    print(f"   {c:6s}: {count:4d} ({pct:5.1f}%)")
print(f"{'='*60}\n")

# ============================================================
# 5️ MODEL FUNCTION (ResNet152)
# ============================================================
def build_resnet152_model(input_shape=(224,224,3), num_classes=3, 
                          dropout=None, regularization=None):
    """ResNet152 with configurable dropout and regularization"""
    base_model = ResNet152(
        weights="imagenet", 
        include_top=False, 
        input_shape=input_shape
    )

    for layer in base_model.layers:
        layer.trainable = True
    
    print(f"   Base model: All {len(base_model.layers)} layers trainable")

    x = base_model.output
    x = GlobalAveragePooling2D(name='global_avg_pool')(x)
    
    # Configure regularization
    reg = None
    if regularization:
        if regularization['type'] == 'l2':
            reg = regularizers.l2(regularization['value'])
        elif regularization['type'] == 'l1':
            reg = regularizers.l1(regularization['value'])
        elif regularization['type'] == 'l1_l2':
            reg = regularizers.l1_l2(l1=regularization['l1'], l2=regularization['l2'])
    
    # Dense layers with optional regularization
    x = Dense(512, activation='relu', kernel_regularizer=reg, name='fc_1')(x)
    if dropout:
        x = Dropout(dropout, name='dropout_1')(x)
    
    x = Dense(256, activation='relu', kernel_regularizer=reg, name='fc_2')(x)
    if dropout:
        x = Dropout(dropout, name='dropout_2')(x)
    
    outputs = Dense(num_classes, activation='softmax', dtype='float32', name='predictions')(x)

    model = Model(inputs=base_model.input, outputs=outputs)
    
    optimizer = AdamW(learning_rate=LEARNING_RATE, weight_decay=1e-4)
    
    model.compile(
        optimizer=optimizer, 
        loss='sparse_categorical_crossentropy',
        metrics=["accuracy"]
    )
    
    print(f"\n   Model Configuration:")
    print(f"   - Dropout: {dropout if dropout else 'None'}")
    print(f"   - Regularization: {regularization if regularization else 'None'}")
    print(f"   - Learning Rate: {LEARNING_RATE}")
    
    return model

# ============================================================
# 6️ DATA GENERATOR
# ============================================================
datagen = ImageDataGenerator(rescale=1./255)

# ============================================================
# 6.5️ LEARNING RATE LOGGER CALLBACK
# ============================================================
class LRLogger(tf.keras.callbacks.Callback):
    """Custom callback to track learning rate in history"""
    def on_epoch_end(self, epoch, logs=None):
        if logs is not None:
            lr = float(tf.keras.backend.get_value(self.model.optimizer.lr))
            logs['lr'] = lr
            if not hasattr(self, 'lrs'):
                self.lrs = []
            self.lrs.append(lr)

# ============================================================
# 7️ GRAD-CAM UTILS
# ============================================================
def generate_gradcam(model, img_array, layer_name):
    """Generate Grad-CAM heatmap"""
    grad_model = Model(inputs=model.inputs, 
                      outputs=[model.get_layer(layer_name).output, model.output])
    with tf.GradientTape() as tape:
        conv_outputs, predictions = grad_model(img_array)
        class_idx = tf.argmax(predictions[0])
        loss = predictions[:, class_idx]
    grads = tape.gradient(loss, conv_outputs)
    pooled_grads = tf.reduce_mean(grads, axis=(0,1,2))
    conv_outputs = conv_outputs[0]
    heatmap = tf.reduce_sum(tf.multiply(pooled_grads, conv_outputs), axis=-1)
    heatmap = tf.maximum(heatmap, 0) / (tf.reduce_max(heatmap) + 1e-10)
    return heatmap.numpy()

def overlay_gradcam(img_path, heatmap, output_path, save_variants=True):
    """Overlay Grad-CAM heatmap on image with multiple visualization variants"""
    try:
        from PIL import Image, ImageFilter
        
        try:
            pil_img = Image.open(img_path).convert('RGB')
        except Exception as e:
            print(f"    Could not read image: {e}")
            return False
        
        pil_img_resized = pil_img.resize((INPUT_SIZE, INPUT_SIZE), Image.LANCZOS)
        img_array = np.array(pil_img_resized)
        
        if heatmap is None or heatmap.size == 0:
            return False
        
        if len(heatmap.shape) != 2:
            return False
        
        heatmap = heatmap.astype(np.float32)
        heatmap_min = heatmap.min()
        heatmap_max = heatmap.max()
        if heatmap_max - heatmap_min > 1e-10:
            heatmap_normalized = (heatmap - heatmap_min) / (heatmap_max - heatmap_min)
        else:
            heatmap_normalized = np.zeros_like(heatmap)
        
        heatmap_pil = Image.fromarray((heatmap_normalized * 255).astype(np.uint8))
        heatmap_resized_pil = heatmap_pil.resize((INPUT_SIZE, INPUT_SIZE), Image.BILINEAR)
        heatmap_resized = np.array(heatmap_resized_pil)
        
        import matplotlib.cm as cm
        colormap = cm.get_cmap('jet')
        heatmap_colored = colormap(heatmap_resized / 255.0)[:, :, :3]
        heatmap_colored = (heatmap_colored * 255).astype(np.uint8)
        
        img_array_float = img_array.astype(np.float32)
        heatmap_colored_float = heatmap_colored.astype(np.float32)
        superimposed_standard = (0.6 * img_array_float + 0.4 * heatmap_colored_float).astype(np.uint8)
        
        result_pil = Image.fromarray(superimposed_standard)
        result_pil.save(output_path)
        
        if save_variants:
            base_path = output_path.replace('.png', '')
            
            pil_img_enhanced = pil_img_resized.filter(ImageFilter.EDGE_ENHANCE_MORE)
            img_enhanced = np.array(pil_img_enhanced).astype(np.float32)
            
            superimposed_embossed = (0.5 * img_enhanced + 0.5 * heatmap_colored_float).astype(np.uint8)
            result_embossed = Image.fromarray(superimposed_embossed)
            result_embossed.save(f"{base_path}_embossed.png")
            
            result_heatmap_only = Image.fromarray(heatmap_colored)
            result_heatmap_only.save(f"{base_path}_heatmap_only.png")
            
            threshold = np.percentile(heatmap_resized, 70)
            mask = (heatmap_resized >= threshold).astype(np.float32)
            mask_3d = np.stack([mask, mask, mask], axis=-1)
            
            superimposed_masked = img_array_float.copy()
            superimposed_masked = (
                (1 - mask_3d) * img_array_float + 
                mask_3d * (0.4 * img_array_float + 0.6 * heatmap_colored_float)
            ).astype(np.uint8)
            
            result_masked = Image.fromarray(superimposed_masked)
            result_masked.save(f"{base_path}_masked.png")
            
            composite_width = INPUT_SIZE * 3
            composite = Image.new('RGB', (composite_width, INPUT_SIZE))
            composite.paste(pil_img_resized, (0, 0))
            composite.paste(result_pil, (INPUT_SIZE, 0))
            composite.paste(result_embossed, (INPUT_SIZE * 2, 0))
            composite.save(f"{base_path}_comparison.png")
        
        return True
            
    except Exception as e:
        print(f"    Grad-CAM overlay error: {e}")
        return False

# ============================================================
# 8️ ROC CURVE PLOTTING
# ============================================================
def plot_roc_curves(y_true, y_pred_proba, class_names, fold_no, save_path):
    """Plot ROC curves for all classes"""
    y_true_bin = label_binarize(y_true, classes=[0, 1, 2])
    n_classes = len(class_names)
    
    fig, ax = plt.subplots(figsize=(10, 8))
    colors = ['blue', 'red', 'green']
    
    for i, color in zip(range(n_classes), colors):
        fpr, tpr, _ = roc_curve(y_true_bin[:, i], y_pred_proba[:, i])
        roc_auc = auc(fpr, tpr)
        ax.plot(fpr, tpr, color=color, lw=2, 
                label=f'{class_names[i]} (AUC = {roc_auc:.3f})')
    
    ax.plot([0, 1], [0, 1], 'k--', lw=2, label='Random (AUC = 0.5)')
    
    ax.set_xlim([0.0, 1.0])
    ax.set_ylim([0.0, 1.05])
    ax.set_xlabel('False Positive Rate', fontsize=12)
    ax.set_ylabel('True Positive Rate', fontsize=12)
    ax.set_title(f'{EXP_ID} - Fold {fold_no} - ROC Curves', fontsize=14)
    ax.legend(loc="lower right", fontsize=10)
    ax.grid(True, alpha=0.3)
    
    plt.tight_layout()
    plt.savefig(save_path, dpi=150)
    plt.close()

# ============================================================
# 9️ WORD/PDF REPORT GENERATION (PER FOLD)
# ============================================================
def generate_fold_report(fold_no, fold_dirs, history, report_df, cm, 
                        bal_acc, overfitting_gap, best_acc,
                        final_train_acc, final_val_acc, y_true, y_pred):
    """Generate comprehensive Word and PDF report for a fold"""
    try:
        doc = Document()
        
        # Title Page
        title = doc.add_heading(f'Lung Cancer Classification - {EXP_ID}', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        subtitle = doc.add_heading(f'Fold {fold_no} Report', level=1)
        subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        doc.add_paragraph()
        
        # Experiment Configuration
        doc.add_heading('1. Experiment Configuration', level=1)
        
        config_table = doc.add_table(rows=8, cols=2)
        config_table.style = 'Light Grid Accent 1'
        
        config_data = [
            ('Experiment ID', EXP_ID),
            ('Experiment Name', exp_config['name']),
            ('Learning Rate', f"{LEARNING_RATE:.0e}"),
            ('Min Learning Rate', f"{MIN_LR:.0e}"),
            ('Dropout', str(DROPOUT) if DROPOUT else 'None'),
            ('Regularization', str(REGULARIZATION) if REGULARIZATION else 'None'),
            ('Batch Size', str(BATCH_SIZE)),
            ('Model', 'ResNet152 (Full Fine-tuning)')
        ]
        
        for i, (key, value) in enumerate(config_data):
            config_table.rows[i].cells[0].text = key
            config_table.rows[i].cells[1].text = str(value)
        
        doc.add_paragraph()
        
        # Performance Summary
        doc.add_heading('2. Performance Summary', level=1)
        
        perf_table = doc.add_table(rows=6, cols=2)
        perf_table.style = 'Light Grid Accent 1'
        
        perf_data = [
            ('Best Validation Accuracy', f"{best_acc:.4f}"),
            ('Final Train Accuracy', f"{final_train_acc:.4f}"),
            ('Final Validation Accuracy', f"{final_val_acc:.4f}"),
            ('Balanced Accuracy', f"{bal_acc:.4f}"),
            ('Overfitting Gap', f"{overfitting_gap:.4f}"),
            ('Epochs Trained', str(len(history.history['loss'])))
        ]
        
        for i, (key, value) in enumerate(perf_data):
            perf_table.rows[i].cells[0].text = key
            perf_table.rows[i].cells[1].text = str(value)
        
        doc.add_paragraph()
        
        # Per-Class Metrics
        doc.add_heading('3. Per-Class Performance', level=1)
        
        class_table = doc.add_table(rows=4, cols=5)
        class_table.style = 'Light Grid Accent 1'
        
        headers = ['Class', 'Precision', 'Recall', 'F1-Score', 'Support']
        for i, header in enumerate(headers):
            class_table.rows[0].cells[i].text = header
        
        for i, cls in enumerate(class_names):
            if cls in report_df.index:
                class_table.rows[i+1].cells[0].text = cls
                class_table.rows[i+1].cells[1].text = f"{report_df.loc[cls, 'precision']:.3f}"
                class_table.rows[i+1].cells[2].text = f"{report_df.loc[cls, 'recall']:.3f}"
                class_table.rows[i+1].cells[3].text = f"{report_df.loc[cls, 'f1-score']:.3f}"
                class_table.rows[i+1].cells[4].text = f"{int(report_df.loc[cls, 'support'])}"
        
        doc.add_paragraph()
        doc.add_page_break()
        
        # Training Curves
        doc.add_heading('4. Training Curves', level=1)
        
        training_curve_path = os.path.join(fold_dirs["plot"], "training_curves.png")
        if os.path.exists(training_curve_path):
            doc.add_picture(training_curve_path, width=Inches(6))
            doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        doc.add_paragraph()
        
        # Learning Rate Schedule
        doc.add_heading('5. Learning Rate Schedule', level=1)
        
        lr_schedule_path = os.path.join(fold_dirs["plot"], "lr_schedule.png")
        if os.path.exists(lr_schedule_path):
            doc.add_picture(lr_schedule_path, width=Inches(6))
            doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        else:
            doc.add_paragraph("Learning rate schedule not available.")
        
        doc.add_paragraph()
        doc.add_page_break()
        
        # Confusion Matrix
        doc.add_heading('6. Confusion Matrix', level=1)
        
        cm_path = os.path.join(fold_dirs["plot"], "confusion_matrix.png")
        if os.path.exists(cm_path):
            doc.add_picture(cm_path, width=Inches(5))
            doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        doc.add_paragraph()
        
        # ROC Curves
        doc.add_heading('7. ROC Curves', level=1)
        
        roc_path = os.path.join(fold_dirs["plot"], "roc_curves.png")
        if os.path.exists(roc_path):
            doc.add_picture(roc_path, width=Inches(6))
            doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        doc.add_paragraph()
        doc.add_page_break()
        
        # Grad-CAM Visualizations
        doc.add_heading('8. Grad-CAM Visualizations', level=1)
        
        doc.add_paragraph('Grad-CAM highlights the regions that the model focuses on for classification.')
        doc.add_paragraph()
        
        for cls_name in class_names:
            doc.add_heading(f'{cls_name} Class - Grad-CAM Examples', level=2)
            
            grad_files = []
            for i in range(1, 4):
                comparison_file = os.path.join(fold_dirs["grad"], 
                                             f"{cls_name}_sample_{i}_gradcam_comparison.png")
                if os.path.exists(comparison_file):
                    grad_files.append(comparison_file)
            
            if grad_files:
                for grad_file in grad_files:
                    doc.add_picture(grad_file, width=Inches(6.5))
                    doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    doc.add_paragraph()
            else:
                doc.add_paragraph(f"No Grad-CAM visualizations available for {cls_name}.")
            
            doc.add_paragraph()
        
        doc.add_page_break()
        
        # Individual Grad-CAM Variants
        doc.add_heading('9. Grad-CAM Variants (Detailed)', level=1)
        
        doc.add_paragraph('Different visualization techniques for better interpretation:')
        variants = [
            ('Standard', 'gradcam.png', 'Standard overlay (60% image + 40% heatmap)'),
            ('Embossed', 'gradcam_embossed.png', 'Edge-enhanced with stronger overlay'),
            ('Heatmap Only', 'gradcam_heatmap_only.png', 'Pure attention map'),
            ('Masked', 'gradcam_masked.png', 'Overlay on high-attention areas only')
        ]
        
        for cls_name in class_names:
            doc.add_heading(f'{cls_name} Class - All Variants', level=2)
            
            sample_base = os.path.join(fold_dirs["grad"], f"{cls_name}_sample_1_")
            
            for variant_name, suffix, description in variants:
                variant_file = sample_base + suffix
                if os.path.exists(variant_file):
                    doc.add_heading(f'{variant_name}', level=3)
                    doc.add_paragraph(description)
                    doc.add_picture(variant_file, width=Inches(4))
                    doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    doc.add_paragraph()
        
        doc.add_page_break()
        
        # Training Log Summary (FIXED SECTION)
        doc.add_heading('10. Training Log Summary', level=1)
        
        total_epochs = len(history.history["loss"])
        doc.add_paragraph(f'Total epochs trained: {total_epochs}')
        doc.add_paragraph()
        
        # Calculate how many epochs to show (first 5 and last 5)
        num_first_epochs = min(5, total_epochs)
        num_last_epochs = min(5, total_epochs)
        
        # Calculate total rows: header + first epochs + separator + last epochs
        total_rows = 1 + num_first_epochs + 1 + num_last_epochs
        
        log_table = doc.add_table(rows=total_rows, cols=6)
        log_table.style = 'Light Grid Accent 1'
        
        # Headers
        log_headers = ['Epoch', 'Loss', 'Accuracy', 'Val Loss', 'Val Acc', 'LR']
        for i, header in enumerate(log_headers):
            log_table.rows[0].cells[i].text = header
        
        # First 5 epochs
        for i in range(num_first_epochs):
            log_table.rows[i+1].cells[0].text = str(i+1)
            log_table.rows[i+1].cells[1].text = f"{history.history['loss'][i]:.4f}"
            log_table.rows[i+1].cells[2].text = f"{history.history['accuracy'][i]:.4f}"
            log_table.rows[i+1].cells[3].text = f"{history.history['val_loss'][i]:.4f}"
            log_table.rows[i+1].cells[4].text = f"{history.history['val_accuracy'][i]:.4f}"
            if 'lr' in history.history:
                log_table.rows[i+1].cells[5].text = f"{history.history['lr'][i]:.2e}"
            else:
                log_table.rows[i+1].cells[5].text = "N/A"
        
        # Separator row (only if there are more than 5 epochs)
        if total_epochs > 5:
            separator_row_idx = num_first_epochs + 1
            log_table.rows[separator_row_idx].cells[0].text = '...'
            for i in range(1, 6):
                log_table.rows[separator_row_idx].cells[i].text = '...'
            
            # Last 5 epochs
            start_idx = max(0, total_epochs - num_last_epochs)
            for i, epoch_idx in enumerate(range(start_idx, total_epochs)):
                row_idx = separator_row_idx + 1 + i
                log_table.rows[row_idx].cells[0].text = str(epoch_idx + 1)
                log_table.rows[row_idx].cells[1].text = f"{history.history['loss'][epoch_idx]:.4f}"
                log_table.rows[row_idx].cells[2].text = f"{history.history['accuracy'][epoch_idx]:.4f}"
                log_table.rows[row_idx].cells[3].text = f"{history.history['val_loss'][epoch_idx]:.4f}"
                log_table.rows[row_idx].cells[4].text = f"{history.history['val_accuracy'][epoch_idx]:.4f}"
                if 'lr' in history.history:
                    log_table.rows[row_idx].cells[5].text = f"{history.history['lr'][epoch_idx]:.2e}"
                else:
                    log_table.rows[row_idx].cells[5].text = "N/A"
        
        # Save Word document
        report_path = os.path.join(fold_dirs["plot"], f"fold_{fold_no}_report.docx")
        doc.save(report_path)
        print(f"    Word report saved: {report_path}")
        
        # Convert to PDF if possible
        if PDF_AVAILABLE:
            try:
                pdf_path = report_path.replace('.docx', '.pdf')
                convert(report_path, pdf_path)
                print(f"    PDF report saved: {pdf_path}")
            except Exception as e:
                print(f"    PDF conversion failed: {e}")
        
    except Exception as e:
        print(f"    Could not generate fold report: {e}")
        import traceback
        traceback.print_exc()
# ============================================================
# 10 COMPREHENSIVE EXPERIMENT REPORT (ALL FOLDS)
# ============================================================
# ============================================================
# 11 COMPREHENSIVE EXPERIMENT REPORT (ALL FOLDS) - FIXED VERSION
# ============================================================
# REPLACE THE ENTIRE generate_experiment_report FUNCTION
# (Original location: approximately lines 703-900)
# with this fixed version:

def generate_experiment_report(output_dir, histories, scores, all_reports, overfitting_gaps):
    """Generate comprehensive Word and PDF report for entire experiment"""
    try:
        doc = Document()
        
        # Title Page
        title = doc.add_heading('Lung Cancer Classification', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        subtitle = doc.add_heading(f'{EXP_ID} - Complete Experiment Report', level=1)
        subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        date_para = doc.add_paragraph(f'Generated: {pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S")}')
        date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        doc.add_page_break()
        
        # Executive Summary
        doc.add_heading('Executive Summary', level=1)
        
        doc.add_paragraph(f'Experiment Name: {exp_config["name"]}')
        doc.add_paragraph(f'Description: {exp_config["description"]}')
        doc.add_paragraph()
        
        summary_table = doc.add_table(rows=6, cols=2)
        summary_table.style = 'Light Grid Accent 1'
        
        summary_data = [
            ('Mean Balanced Accuracy', f"{np.mean(scores):.4f} ± {np.std(scores):.4f}"),
            ('Mean Overfitting Gap', f"{np.mean(overfitting_gaps):.4f} ± {np.std(overfitting_gaps):.4f}"),
            ('Best Fold Performance', f"{max(scores):.4f} (Fold {np.argmax(scores)+1})"),
            ('Worst Fold Performance', f"{min(scores):.4f} (Fold {np.argmin(scores)+1})"),
            ('Average Epochs Trained', f"{np.mean([len(h.history['loss']) for h in histories]):.1f}"),
            ('Total Training Time', 'See individual fold reports')
        ]
        
        for i, (key, value) in enumerate(summary_data):
            summary_table.rows[i].cells[0].text = key
            summary_table.rows[i].cells[1].text = str(value)
        
        doc.add_page_break()
        
        # Experiment Configuration
        doc.add_heading('Experiment Configuration', level=1)
        
        config_table = doc.add_table(rows=10, cols=2)
        config_table.style = 'Light Grid Accent 1'
        
        config_data = [
            ('Experiment ID', EXP_ID),
            ('Experiment Name', exp_config['name']),
            ('Model Architecture', 'ResNet152 (ImageNet pretrained)'),
            ('Training Strategy', 'Full fine-tuning (all layers trainable)'),
            ('Learning Rate', f"{LEARNING_RATE:.0e}"),
            ('Min Learning Rate', f"{MIN_LR:.0e}"),
            ('Dropout', str(DROPOUT) if DROPOUT else 'None'),
            ('Regularization', str(REGULARIZATION) if REGULARIZATION else 'None'),
            ('Batch Size', str(BATCH_SIZE)),
            ('Input Size', f'{INPUT_SIZE}x{INPUT_SIZE}')
        ]
        
        for i, (key, value) in enumerate(config_data):
            config_table.rows[i].cells[0].text = key
            config_table.rows[i].cells[1].text = str(value)
        
        doc.add_page_break()
        
        # Cross-Validation Results
        doc.add_heading('Cross-Validation Results (5-Fold)', level=1)
        
        cv_table = doc.add_table(rows=6, cols=6)
        cv_table.style = 'Light Grid Accent 1'
        
        cv_headers = ['Fold', 'Balanced Acc', 'Val Acc', 'Overfit Gap', 'Epochs', 'Status']
        for i, header in enumerate(cv_headers):
            cv_table.rows[0].cells[i].text = header
        
        for i in range(5):
            cv_table.rows[i+1].cells[0].text = str(i+1)
            cv_table.rows[i+1].cells[1].text = f"{scores[i]:.4f}"
            cv_table.rows[i+1].cells[2].text = f"{histories[i].history['val_accuracy'][-1]:.4f}"
            cv_table.rows[i+1].cells[3].text = f"{overfitting_gaps[i]:.4f}"
            cv_table.rows[i+1].cells[4].text = str(len(histories[i].history['loss']))
            
            if scores[i] == max(scores):
                cv_table.rows[i+1].cells[5].text = ' Best'
            elif scores[i] == min(scores):
                cv_table.rows[i+1].cells[5].text = ' Worst'
            else:
                cv_table.rows[i+1].cells[5].text = '✓'
        
        doc.add_paragraph()
        
        # Average Per-Class Performance
        doc.add_heading('Average Per-Class Performance', level=1)
        
        avg_report = pd.concat(all_reports).groupby(level=0).mean()
        
        class_perf_table = doc.add_table(rows=4, cols=4)
        class_perf_table.style = 'Light Grid Accent 1'
        
        class_headers = ['Class', 'Precision', 'Recall', 'F1-Score']
        for i, header in enumerate(class_headers):
            class_perf_table.rows[0].cells[i].text = header
        
        for i, cls in enumerate(class_names):
            if cls in avg_report.index:
                class_perf_table.rows[i+1].cells[0].text = cls
                class_perf_table.rows[i+1].cells[1].text = f"{avg_report.loc[cls, 'precision']:.3f}"
                class_perf_table.rows[i+1].cells[2].text = f"{avg_report.loc[cls, 'recall']:.3f}"
                class_perf_table.rows[i+1].cells[3].text = f"{avg_report.loc[cls, 'f1-score']:.3f}"
        
        doc.add_page_break()
        
        # Fold-by-Fold Summary (SIMPLIFIED - No detailed tables for each fold)
        for fold_idx in range(5):
            doc.add_heading(f'Fold {fold_idx+1} - Summary', level=1)
            
            # Simple metrics list instead of table
            doc.add_paragraph(f"Balanced Accuracy: {scores[fold_idx]:.4f}")
            doc.add_paragraph(f"Validation Accuracy: {histories[fold_idx].history['val_accuracy'][-1]:.4f}")
            doc.add_paragraph(f"Overfitting Gap: {overfitting_gaps[fold_idx]:.4f}")
            doc.add_paragraph(f"Epochs Trained: {len(histories[fold_idx].history['loss'])}")
            doc.add_paragraph()
            
            # Training curves
            fold_plot_dir = os.path.join(plot_dir, f"fold_{fold_idx+1}")
            training_curve = os.path.join(fold_plot_dir, "training_curves.png")
            if os.path.exists(training_curve):
                doc.add_picture(training_curve, width=Inches(6))
                doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            doc.add_paragraph()
            
            # Confusion matrix
            cm_path = os.path.join(fold_plot_dir, "confusion_matrix.png")
            if os.path.exists(cm_path):
                doc.add_picture(cm_path, width=Inches(4.5))
                doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            if fold_idx < 4:
                doc.add_page_break()
        
        # Conclusions
        doc.add_page_break()
        doc.add_heading('Conclusions', level=1)
        
        doc.add_paragraph('Key Findings:')
        findings = [
            f"The model achieved a mean balanced accuracy of {np.mean(scores):.4f} across 5 folds.",
            f"The mean overfitting gap was {np.mean(overfitting_gaps):.4f}, indicating {'good generalization' if np.mean(overfitting_gaps) < 0.1 else 'potential overfitting'}.",
            f"Best performing fold: Fold {np.argmax(scores)+1} with {max(scores):.4f} balanced accuracy.",
            f"Most consistent performance with standard deviation of {np.std(scores):.4f}.",
        ]
        
        for finding in findings:
            p = doc.add_paragraph(finding, style='List Bullet')
        
        doc.add_paragraph()
        doc.add_paragraph('All detailed results, models, and visualizations are saved in:')
        doc.add_paragraph(f'{output_dir}', style='Quote')
        
        # Save Word document
        report_path = os.path.join(output_dir, f'{EXP_ID}_complete_report.docx')
        doc.save(report_path)
        print(f"\n Comprehensive Word report saved: {report_path}")
        
        # Convert to PDF if possible
        if PDF_AVAILABLE:
            try:
                pdf_path = report_path.replace('.docx', '.pdf')
                convert(report_path, pdf_path)
                print(f" Comprehensive PDF report saved: {pdf_path}")
            except Exception as e:
                print(f" PDF conversion failed: {e}")
        
    except Exception as e:
        print(f" Could not generate comprehensive report: {e}")
        import traceback
        traceback.print_exc()
# ============================================================
# 1️ 1 K-FOLD CROSS VALIDATION (MAIN TRAINING LOOP)
# ============================================================
kf = KFold(n_splits=5, shuffle=True, random_state=SEED)
fold_no = 1
histories, scores = [], []
all_reports = []

for train_idx, val_idx in kf.split(image_paths):
    print(f"\n{'='*60}")
    print(f" {EXP_ID} - TRAINING FOLD {fold_no}/5")
    print(f"{'='*60}")

    tf.keras.backend.clear_session()
    
    fold_dirs = {
        "model": os.path.join(model_dir, f"fold_{fold_no}"),
        "plot": os.path.join(plot_dir, f"fold_{fold_no}"),
        "grad": os.path.join(gradcam_dir, f"fold_{fold_no}")
    }
    for d in fold_dirs.values(): 
        os.makedirs(d, exist_ok=True)
    log_file = os.path.join(log_dir, f"fold_{fold_no}_log.csv")

    train_df = pd.DataFrame({
        'filename': image_paths[train_idx], 
        'class': [class_names[l] for l in labels[train_idx]]
    })
    val_df = pd.DataFrame({
        'filename': image_paths[val_idx], 
        'class': [class_names[l] for l in labels[val_idx]]
    })

    print("\n Class distribution:")
    print("Training set:")
    for cls_name in class_names:
        count = np.sum(train_df['class'] == cls_name)
        print(f"  {cls_name:6s}: {count:4d}")
    print("Validation set:")
    for cls_name in class_names:
        count = np.sum(val_df['class'] == cls_name)
        print(f"  {cls_name:6s}: {count:4d}")

    train_gen = datagen.flow_from_dataframe(
        train_df, x_col='filename', y_col='class',
        target_size=(INPUT_SIZE, INPUT_SIZE), 
        class_mode='sparse',
        classes=class_names,
        batch_size=BATCH_SIZE, shuffle=True
    )
    val_gen = datagen.flow_from_dataframe(
        val_df, x_col='filename', y_col='class',
        target_size=(INPUT_SIZE, INPUT_SIZE), 
        class_mode='sparse',
        classes=class_names,
        batch_size=BATCH_SIZE, shuffle=False
    )

    model = build_resnet152_model(dropout=DROPOUT, regularization=REGULARIZATION)

    lr_logger = LRLogger()
    
    callbacks = [
        lr_logger,
        CSVLogger(log_file, append=True),
        ReduceLROnPlateau(
            monitor='val_loss', 
            factor=0.5,
            patience=7,
            min_lr=MIN_LR,
            verbose=1
        ),
        EarlyStopping(
            monitor='val_loss', 
            patience=15,
            restore_best_weights=True, 
            verbose=1
        )
    ]

    print(f"\n Starting training...")
    
    history = model.fit(
        train_gen,
        validation_data=val_gen,
        epochs=1000,
        callbacks=callbacks,
        verbose=1
    )

    print(f"\n Evaluating Fold {fold_no}...")
    val_gen.reset()
    y_pred_proba = model.predict(val_gen, verbose=0)
    y_pred = np.argmax(y_pred_proba, axis=1)
    
    y_true = np.array([class_names.index(label) for label in val_df['class'].values])
    
    best_acc = max(history.history['val_accuracy'])
    final_train_acc = history.history['accuracy'][-1]
    final_val_acc = history.history['val_accuracy'][-1]
    bal_acc = balanced_accuracy_score(y_true, y_pred)
    
    overfitting_gap = final_train_acc - final_val_acc
    
    print(f"\n Fold {fold_no} Results:")
    print(f"   Best Val Accuracy: {best_acc:.4f}")
    print(f"   Final Train Accuracy: {final_train_acc:.4f}")
    print(f"   Final Val Accuracy: {final_val_acc:.4f}")
    print(f"   Balanced Accuracy: {bal_acc:.4f}")
    print(f"   Overfitting Gap: {overfitting_gap:.4f}")

    histories.append(history)
    scores.append(bal_acc)

    report = classification_report(y_true, y_pred, target_names=class_names, 
                                   output_dict=True, zero_division=0)
    report_df = pd.DataFrame(report).transpose()
    all_reports.append(report_df)
    
    print(f"\n Per-class metrics:")
    for cls in class_names:
        if cls in report:
            print(f"  {cls:6s}: Precision={report[cls]['precision']:.3f}, "
                  f"Recall={report[cls]['recall']:.3f}, "
                  f"F1={report[cls]['f1-score']:.3f}")
    
    report_df.to_csv(os.path.join(fold_dirs["plot"], f"class_report.csv"))

    model_path = os.path.join(fold_dirs["model"], f"resnet152_fold_{fold_no}.keras")
    model.save(model_path, save_format='keras')
    print(f" Complete model saved: {model_path}")
    
    weights_path = os.path.join(fold_dirs["model"], f"resnet152_fold_{fold_no}_weights.h5")
    model.save_weights(weights_path)
    print(f" Weights saved: {weights_path}")

    # Confusion Matrix
    cm = confusion_matrix(y_true, y_pred)
    fig, ax = plt.subplots(figsize=(8, 6))
    disp = ConfusionMatrixDisplay(confusion_matrix=cm, display_labels=class_names)
    disp.plot(cmap=plt.cm.Blues, ax=ax)
    plt.title(f"{EXP_ID} - Fold {fold_no}\nBalanced Accuracy: {bal_acc:.3f} | Overfitting Gap: {overfitting_gap:.3f}")
    plt.tight_layout()
    plt.savefig(os.path.join(fold_dirs["plot"], f"confusion_matrix.png"), dpi=150)
    plt.close()

    # ROC Curves
    plot_roc_curves(y_true, y_pred_proba, class_names, fold_no,
                   os.path.join(fold_dirs["plot"], f"roc_curves.png"))

    # Training Curves
    fig, axes = plt.subplots(1, 2, figsize=(14, 5))
    
    epochs = range(1, len(history.history['accuracy']) + 1)
    
    axes[0].plot(epochs, history.history['accuracy'], label='Train', linewidth=2, alpha=0.8)
    axes[0].plot(epochs, history.history['val_accuracy'], label='Validation', linewidth=2, alpha=0.8)
    axes[0].fill_between(epochs, history.history['accuracy'], history.history['val_accuracy'], 
                         alpha=0.2, color='red', label='Overfitting Gap')
    axes[0].set_title(f'{EXP_ID} - Fold {fold_no} - Accuracy')
    axes[0].set_xlabel('Epochs')
    axes[0].set_ylabel('Accuracy')
    axes[0].legend()
    axes[0].grid(True, alpha=0.3)
    
    axes[1].plot(epochs, history.history['loss'], label='Train', linewidth=2, alpha=0.8)
    axes[1].plot(epochs, history.history['val_loss'], label='Validation', linewidth=2, alpha=0.8)
    axes[1].set_title(f'{EXP_ID} - Fold {fold_no} - Loss')
    axes[1].set_xlabel('Epochs')
    axes[1].set_ylabel('Loss')
    axes[1].legend()
    axes[1].grid(True, alpha=0.3)
    
    plt.tight_layout()
    plt.savefig(os.path.join(fold_dirs["plot"], f"training_curves.png"), dpi=150)
    plt.close()

    # Learning Rate Schedule Plot
    if hasattr(lr_logger, 'lrs') and len(lr_logger.lrs) > 0:
        lr_values = lr_logger.lrs
    elif 'lr' in history.history:
        lr_values = history.history['lr']
    else:
        lr_values = None
    
    if lr_values is not None:
        fig, ax = plt.subplots(figsize=(10, 5))
        epochs_range = range(1, len(lr_values) + 1)
        ax.plot(epochs_range, lr_values, linewidth=2, color='purple', marker='o', 
                markersize=3, alpha=0.7)
        ax.set_title(f'{EXP_ID} - Fold {fold_no} - Learning Rate Schedule', fontsize=14)
        ax.set_xlabel('Epochs', fontsize=12)
        ax.set_ylabel('Learning Rate', fontsize=12)
        ax.set_yscale('log')
        ax.grid(True, alpha=0.3, which='both')
        ax.axhline(y=MIN_LR, color='r', linestyle='--', linewidth=1.5, 
                   label=f'Min LR = {MIN_LR:.0e}')
        ax.axhline(y=LEARNING_RATE, color='g', linestyle='--', linewidth=1.5, 
                   label=f'Initial LR = {LEARNING_RATE:.0e}')
        ax.legend(fontsize=10)
        
        for i in range(1, len(lr_values)):
            if lr_values[i] < lr_values[i-1]:
                ax.annotate('LR reduced', xy=(i+1, lr_values[i]), 
                           xytext=(10, 10), textcoords='offset points',
                           fontsize=8, color='red', alpha=0.7,
                           arrowprops=dict(arrowstyle='->', color='red', alpha=0.5))
        
        plt.tight_layout()
        plt.savefig(os.path.join(fold_dirs["plot"], f"lr_schedule.png"), dpi=150)
        plt.close()
        print(f"    Learning rate schedule saved")

    # Grad-CAM Generation
    print(f"\n Generating Grad-CAM visualizations...")
    conv_layer_name = 'conv5_block3_out'
    
    n_samples = 3
    successful_gradcams = 0
    total_attempts = 0
    
    for cls_idx, cls_name in enumerate(class_names):
        correct_mask = (y_true == cls_idx) & (y_pred == cls_idx)
        correct_indices = np.where(correct_mask)[0]
        
        if len(correct_indices) == 0:
            print(f"    No correct predictions for {cls_name}")
            continue
        
        sample_indices = np.random.choice(correct_indices, 
                                         min(n_samples, len(correct_indices)), 
                                         replace=False)
        
        for i, idx in enumerate(sample_indices):
            total_attempts += 1
            img_path = val_df.iloc[idx]['filename']
            
            try:
                img = tf.keras.preprocessing.image.load_img(img_path, 
                                                            target_size=(INPUT_SIZE, INPUT_SIZE))
                img_array = tf.keras.preprocessing.image.img_to_array(img)
                img_array = np.expand_dims(img_array, axis=0) / 255.0
            except Exception as e:
                print(f"   ✗ Could not load {cls_name} sample {i+1}: {e}")
                continue
            
            try:
                heatmap = generate_gradcam(model, img_array, conv_layer_name)
            except Exception as e:
                print(f"   ✗ Heatmap generation failed for {cls_name} sample {i+1}: {e}")
                continue
            
            output_path = os.path.join(fold_dirs["grad"], 
                                      f"{cls_name}_sample_{i+1}_gradcam.png")
            success = overlay_gradcam(img_path, heatmap, output_path, save_variants=True)
            
            if success:
                successful_gradcams += 1
                print(f"   ✓ Generated Grad-CAM variants for {cls_name} sample {i+1}")
    
    print(f"    Grad-CAM success rate: {successful_gradcams}/{total_attempts}")
    print(f"    Total images generated: {successful_gradcams * 5} (5 variants per sample)")

    # Generate Word/PDF Report for this fold
    if DOCX_AVAILABLE:
        print(f"\n Generating reports for Fold {fold_no}...")
        generate_fold_report(fold_no, fold_dirs, history, report_df, cm, 
                            bal_acc, overfitting_gap, best_acc, 
                            final_train_acc, final_val_acc, y_true, y_pred)
    else:
        print(f"\n Skipping report generation (python-docx not installed)")

    fold_no += 1

# ============================================================
# 1️ 2️ FINAL SUMMARY & REPORT GENERATION
# ============================================================
print(f"\n{'='*60}")
print(f" CROSS-VALIDATION SUMMARY - {EXP_ID}")
print(f"{'='*60}")
print(f"\n Experiment: {exp_config['name']}")
print(f" Description: {exp_config['description']}")
print(f" Learning Rate: {LEARNING_RATE}")
print(f" Dropout: {DROPOUT if DROPOUT else 'None'}")
print(f" Regularization: {REGULARIZATION if REGULARIZATION else 'None'}")

print("\n Balanced Accuracy per fold:")
for i, s in enumerate(scores, 1):
    print(f"  Fold {i}: {s:.4f}")
print(f"\n Mean Balanced Accuracy: {np.mean(scores):.4f} ± {np.std(scores):.4f}")

overfitting_gaps = []
for h in histories:
    gap = h.history['accuracy'][-1] - h.history['val_accuracy'][-1]
    overfitting_gaps.append(gap)
print(f" Mean Overfitting Gap: {np.mean(overfitting_gaps):.4f} ± {np.std(overfitting_gaps):.4f}")

print(f"\n Average Per-Class Metrics:")
avg_report = pd.concat(all_reports).groupby(level=0).mean()
for cls in class_names:
    if cls in avg_report.index:
        print(f"  {cls:6s}: Precision={avg_report.loc[cls, 'precision']:.3f}, "
              f"Recall={avg_report.loc[cls, 'recall']:.3f}, "
              f"F1={avg_report.loc[cls, 'f1-score']:.3f}")

print(f"\n All results saved to: {output_dir}")
print(f"{'='*60}\n")

# Save comprehensive summary
summary_log = {
    'experiment_id': [EXP_ID] * 5,
    'experiment_name': [exp_config['name']] * 5,
    'learning_rate': [LEARNING_RATE] * 5,
    'min_lr': [MIN_LR] * 5,
    'dropout': [DROPOUT] * 5,
    'regularization': [str(REGULARIZATION)] * 5,
    'fold': list(range(1, 6)),
    'balanced_accuracy': scores,
    'best_val_accuracy': [max(h.history['val_accuracy']) for h in histories],
    'final_train_accuracy': [h.history['accuracy'][-1] for h in histories],
    'final_val_accuracy': [h.history['val_accuracy'][-1] for h in histories],
    'overfitting_gap': overfitting_gaps,
    'final_train_loss': [h.history['loss'][-1] for h in histories],
    'final_val_loss': [h.history['val_loss'][-1] for h in histories],
    'epochs_trained': [len(h.history['loss']) for h in histories]
}
summary_df = pd.DataFrame(summary_log)
summary_df.to_csv(os.path.join(log_dir, 'training_summary.csv'), index=False)

avg_report.to_csv(os.path.join(log_dir, 'average_class_report.csv'))

# Create experiment comparison summary
comparison_summary = {
    'Experiment_ID': EXP_ID,
    'Experiment_Name': exp_config['name'],
    'Learning_Rate': LEARNING_RATE,
    'Min_LR': MIN_LR,
    'Dropout': DROPOUT if DROPOUT else 'None',
    'Regularization': str(REGULARIZATION) if REGULARIZATION else 'None',
    'Mean_Balanced_Accuracy': np.mean(scores),
    'Std_Balanced_Accuracy': np.std(scores),
    'Mean_Overfitting_Gap': np.mean(overfitting_gaps),
    'Std_Overfitting_Gap': np.std(overfitting_gaps),
    'Mean_Best_Val_Acc': np.mean([max(h.history['val_accuracy']) for h in histories]),
    'Mean_Final_Train_Acc': np.mean([h.history['accuracy'][-1] for h in histories]),
    'Mean_Final_Val_Acc': np.mean([h.history['val_accuracy'][-1] for h in histories]),
    'Mean_Epochs_Trained': np.mean([len(h.history['loss']) for h in histories]),
}

comparison_file = os.path.join("/data/vaishnav25/ADITHYA", "all_experiments_comparison.csv")
comparison_df = pd.DataFrame([comparison_summary])

if os.path.exists(comparison_file):
    existing_df = pd.read_csv(comparison_file)
    if EXP_ID in existing_df['Experiment_ID'].values:
        existing_df = existing_df[existing_df['Experiment_ID'] != EXP_ID]
    comparison_df = pd.concat([existing_df, comparison_df], ignore_index=True)

comparison_df.to_csv(comparison_file, index=False)
print(f"\n📊 Experiment comparison saved to: {comparison_file}")

# Detailed text report
with open(os.path.join(log_dir, 'final_report.txt'), 'w') as f:
    f.write("="*70 + "\n")
    f.write(f"LUNG CANCER CLASSIFICATION - {EXP_ID} REPORT\n")
    f.write("="*70 + "\n\n")
    
    f.write("EXPERIMENT CONFIGURATION:\n")
    f.write(f"  Experiment ID: {EXP_ID}\n")
    f.write(f"  Name: {exp_config['name']}\n")
    f.write(f"  Description: {exp_config['description']}\n")
    f.write(f"  Learning Rate: {LEARNING_RATE}\n")
    f.write(f"  Min Learning Rate: {MIN_LR}\n")
    f.write(f"  Dropout: {DROPOUT if DROPOUT else 'None'}\n")
    f.write(f"  Regularization: {REGULARIZATION if REGULARIZATION else 'None'}\n\n")
    
    f.write("CROSS-VALIDATION RESULTS:\n")
    f.write(f"Mean Balanced Accuracy: {np.mean(scores):.4f} ± {np.std(scores):.4f}\n")
    f.write(f"Mean Overfitting Gap: {np.mean(overfitting_gaps):.4f} ± {np.std(overfitting_gaps):.4f}\n\n")
    
    f.write("Per-fold results:\n")
    for i in range(5):
        f.write(f"  Fold {i+1}:\n")
        f.write(f"    Balanced Accuracy: {scores[i]:.4f}\n")
        f.write(f"    Overfitting Gap: {overfitting_gaps[i]:.4f}\n")
        f.write(f"    Final Train Acc: {summary_df.loc[i, 'final_train_accuracy']:.4f}\n")
        f.write(f"    Final Val Acc: {summary_df.loc[i, 'final_val_accuracy']:.4f}\n")
        f.write(f"    Epochs Trained: {summary_df.loc[i, 'epochs_trained']}\n\n")

# ============================================================
# 1️ 3️ GENERATE COMPREHENSIVE EXPERIMENT REPORT (CRITICAL)
# ============================================================
if DOCX_AVAILABLE:
    print(f"\n{'='*60}")
    print(f" GENERATING COMPREHENSIVE EXPERIMENT REPORT")
    print(f"{'='*60}")
    generate_experiment_report(output_dir, histories, scores, all_reports, overfitting_gaps)
    print(f"\n Complete experiment report (all folds) generated!")
    print(f"    Location: {output_dir}/{EXP_ID}_complete_report.docx")
    if PDF_AVAILABLE:
        print(f"    PDF: {output_dir}/{EXP_ID}_complete_report.pdf")
else:
    print(f"\n Skipping comprehensive report (python-docx not installed)")

print(f"\n {EXP_ID} Training Complete!")

# ============================================================
# 1️ 4️ FINAL SUMMARY OUTPUT
# ============================================================
print(f"\n{'='*70}")
print(" REPORT GENERATION SUMMARY:")
print("="*70)
print(f"\n Per-Fold Reports (5 reports):")
print(f"   Location: {output_dir}/plots/fold_{{1-5}}/")
print(f"   Files:")
print(f"     - fold_{{1-5}}_report.docx")
if PDF_AVAILABLE:
    print(f"     - fold_{{1-5}}_report.pdf")

print(f"\n Comprehensive Experiment Report (1 report):")
print(f"   Location: {output_dir}/")
print(f"   Files:")
print(f"     - {EXP_ID}_complete_report.docx")
if PDF_AVAILABLE:
    print(f"     - {EXP_ID}_complete_report.pdf")

print(f"\n Other Outputs:")
print(f"   - Models: {output_dir}/models/fold_{{1-5}}/")
print(f"   - Plots: {output_dir}/plots/fold_{{1-5}}/")
print(f"   - Grad-CAMs: {output_dir}/gradcams/fold_{{1-5}}/")
print(f"   - Logs: {output_dir}/logs/")

print(f"\n{'='*70}")
print(" AVAILABLE EXPERIMENTS (14 total):")
print("="*70)
print("\n Learning Rate Experiments (6):")
for exp_id in ['EXP_1', 'EXP_2', 'EXP_3', 'EXP_4', 'EXP_5', 'EXP_6']:
    print(f"  {exp_id}: {EXPERIMENTS[exp_id]['description']}")

print("\n Dropout Experiments (4):")
for exp_id in ['EXP_7', 'EXP_8', 'EXP_9', 'EXP_10']:
    print(f"  {exp_id}: {EXPERIMENTS[exp_id]['description']}")

print("\n Regularization Experiments (4):")
for exp_id in ['EXP_11', 'EXP_12', 'EXP_13', 'EXP_14']:
    print(f"  {exp_id}: {EXPERIMENTS[exp_id]['description']}")

print(f"\n To run a different experiment, change EXP_ID at line 83")
print(f" Track all experiments in: {comparison_file}")
print(f"{'='*70}\n")

# Display comparison table if multiple experiments completed
if os.path.exists(comparison_file):
    all_exps = pd.read_csv(comparison_file)
    if len(all_exps) > 1:
        print(f"\n{'='*70}")
        print(" COMPARISON ACROSS ALL COMPLETED EXPERIMENTS:")
        print(f"{'='*70}")
        
        all_exps_sorted = all_exps.sort_values('Mean_Balanced_Accuracy', ascending=False)
        
        print(f"\n{'Rank':<6} {'Exp_ID':<10} {'Name':<30} {'Bal_Acc':<10} {'Overfit':<10}")
        print("-" * 70)
        for idx, row in all_exps_sorted.iterrows():
            rank = list(all_exps_sorted.index).index(idx) + 1
            print(f"{rank:<6} {row['Experiment_ID']:<10} {row['Experiment_Name']:<30} "
                  f"{row['Mean_Balanced_Accuracy']:.4f}    {row['Mean_Overfitting_Gap']:.4f}")
        
        best_exp = all_exps_sorted.iloc[0]
        print(f"\n Best Experiment: {best_exp['Experiment_ID']} - {best_exp['Experiment_Name']}")
        print(f"   Balanced Accuracy: {best_exp['Mean_Balanced_Accuracy']:.4f} ± {best_exp['Std_Balanced_Accuracy']:.4f}")
        print(f"   Overfitting Gap: {best_exp['Mean_Overfitting_Gap']:.4f}")
        print(f"{'='*70}\n")

print(f"\n{'='*70}")
print(" ALL PROCESSING COMPLETE!")
print(f"{'='*70}")
print(f"\nResults saved to: {output_dir}")
print(f"Comparison file: {comparison_file}")
print(f"\n Experiment {EXP_ID} successfully completed with full reporting!")
print(f"{'='*70}\n")

# ============================================================
# 1️ 5️ ZIP THE OUTPUT FOLDER
# ============================================================
print(f"\n{'='*70}")
print(" CREATING ZIP ARCHIVE")
print(f"{'='*70}")

import shutil
from datetime import datetime

try:
    # Create zip filename with timestamp
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    zip_base_name = f"RESNET152_{EXP_ID}_{timestamp}"
    zip_path = os.path.join("/data/vaishnav25/ADITHYA", zip_base_name)
    
    print(f"\n Compressing experiment folder...")
    print(f"   Source: {output_dir}")
    print(f"   Target: {zip_path}.zip")
    
    # Create zip archive (this will create {zip_path}.zip)
    shutil.make_archive(zip_path, 'zip', output_dir)
    
    # Get zip file size
    zip_file_path = f"{zip_path}.zip"
    zip_size_bytes = os.path.getsize(zip_file_path)
    zip_size_mb = zip_size_bytes / (1024 * 1024)
    zip_size_gb = zip_size_bytes / (1024 * 1024 * 1024)
    
    if zip_size_gb >= 1:
        size_str = f"{zip_size_gb:.2f} GB"
    else:
        size_str = f"{zip_size_mb:.2f} MB"
    
    print(f"\n ZIP archive created successfully!")
    print(f"    Location: {zip_file_path}")
    print(f"    Size: {size_str}")
    print(f"    Contains: All fold results, models, reports, and visualizations")
    
    # Optional: Create a manifest file inside the archive info
    manifest_path = os.path.join(output_dir, "archive_manifest.txt")
    with open(manifest_path, 'w') as f:
        f.write(f"{'='*70}\n")
        f.write(f"EXPERIMENT ARCHIVE MANIFEST\n")
        f.write(f"{'='*70}\n\n")
        f.write(f"Archive created: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        f.write(f"Experiment ID: {EXP_ID}\n")
        f.write(f"Experiment Name: {exp_config['name']}\n")
        f.write(f"Archive Size: {size_str}\n\n")
        
        f.write(f"CONTENTS:\n")
        f.write(f"  - 5 Fold Reports (Word + PDF)\n")
        f.write(f"  - 1 Comprehensive Report (Word + PDF)\n")
        f.write(f"  - 5 Trained Models (.keras + .h5)\n")
        f.write(f"  - Training curves, confusion matrices, ROC curves\n")
        f.write(f"  - Learning rate schedules\n")
        f.write(f"  - Grad-CAM visualizations (5 variants × 3 samples × 3 classes)\n")
        f.write(f"  - CSV logs and metrics\n")
        f.write(f"  - Augmented Poor class images\n\n")
        
        f.write(f"RESULTS SUMMARY:\n")
        f.write(f"  Mean Balanced Accuracy: {np.mean(scores):.4f} ± {np.std(scores):.4f}\n")
        f.write(f"  Mean Overfitting Gap: {np.mean(overfitting_gaps):.4f} ± {np.std(overfitting_gaps):.4f}\n")
        f.write(f"  Total Epochs Trained: {sum([len(h.history['loss']) for h in histories])}\n")
        f.write(f"  Best Fold: Fold {np.argmax(scores)+1} ({max(scores):.4f})\n\n")
        
        f.write(f"TO EXTRACT:\n")
        f.write(f"  unzip {os.path.basename(zip_file_path)}\n")
        f.write(f"{'='*70}\n")
    
    print(f"\n Archive manifest created: {manifest_path}")
    
    # Optional: Keep or delete original folder (commented out for safety)
    # print(f"\n Note: Original folder retained at: {output_dir}")
    # Uncomment below to delete original folder after zipping:
    # shutil.rmtree(output_dir)
    # print(f"   Original folder deleted (contents preserved in ZIP)")
    
except Exception as e:
    print(f"\n Error creating ZIP archive: {e}")
    print(f"   Original folder preserved at: {output_dir}")
    import traceback
    traceback.print_exc()

print(f"\n{'='*70}")
print(" EXPERIMENT COMPLETE WITH ARCHIVING!")
print(f"{'='*70}")
print(f"\n ZIP Archive: {zip_file_path if 'zip_file_path' in locals() else 'Not created'}")
print(f" Original Folder: {output_dir}")
print(f" Comparison File: {comparison_file}")
print(f"\n All done! Experiment {EXP_ID} completed and archived successfully!")
print(f"{'='*70}\n")